from __future__ import annotations

import io
import re
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


class DOCXRenderer:
    """
    Converts plain-text document output from Claude into a .docx file (in-memory bytes).
    python-docx is used for reliable, dependency-free Word generation.
    """

    def render(self, document_text: str, collected_data: Dict[str, Any] | None = None) -> bytes:
        doc = Document()

        # --- Page setup (A4) ---
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)

        # --- Split text into paragraphs ---
        paragraphs = document_text.split("\n\n")
        first = True

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            if first and (
                "LAST WILL" in para_text or "TESTAMENT" in para_text
            ):
                # Title
                heading = doc.add_heading(para_text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first = False
                continue

            first = False

            # Section number heading (e.g. "1. REVOCATION")
            if re.match(r"^\d+\.\s+[A-Z]", para_text):
                p = doc.add_paragraph()
                run = p.add_run(para_text.replace("\n", " "))
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                continue

            # NOTE section
            if para_text.startswith("---"):
                p = doc.add_paragraph(para_text[3:].strip().replace("\n", " "))
                p.style.font.size = Pt(10)
                continue

            # Regular paragraph
            p = doc.add_paragraph(para_text.replace("\n", " "))
            run = p.runs[0] if p.runs else p.add_run("")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # --- Serialise to bytes ---
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
